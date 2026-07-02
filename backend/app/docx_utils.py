import io

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from an uploaded .docx file, preserving paragraph
    breaks. Tables are read row by row as tab-separated text."""
    docx = DocxDocument(io.BytesIO(file_bytes))
    parts = []
    for para in docx.paragraphs:
        parts.append(para.text)
    for table in docx.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip()


def build_docx_from_text(title: str, content: str) -> bytes:
    """Build a .docx file from plain text content, one paragraph per line."""
    docx = DocxDocument()
    docx.add_heading(title, level=1)
    for line in content.split("\n"):
        docx.add_paragraph(line)
    buffer = io.BytesIO()
    docx.save(buffer)
    return buffer.getvalue()


INVOICE_STATUS_LABEL = {"draft": "טיוטה", "sent": "נשלחה", "paid": "שולמה"}


def build_invoice_docx(invoice, case, client) -> bytes:
    """Build a client-ready Word document for an already-generated Invoice,
    itemizing the time entries snapshotted into it (see routers/invoices.py -
    every entry attached to an invoice is by construction billable and
    priced, since that's a precondition of invoice generation)."""
    docx = DocxDocument()

    docx.add_heading("חשבונית שירותים משפטיים", level=1)

    meta = docx.add_paragraph()
    meta.add_run(f"מספר חשבונית: {invoice.invoice_number}\n")
    meta.add_run(f"לקוח: {client.full_name}\n")
    meta.add_run(f"תיק: {case.title} ({case.case_number})\n")
    meta.add_run(f"תאריך הפקה: {invoice.issue_date.strftime('%d/%m/%Y')}\n")
    meta.add_run(f"סטטוס: {INVOICE_STATUS_LABEL.get(invoice.status.value, invoice.status.value)}")

    table = docx.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    header_cells = table.rows[0].cells
    headers = ["תאריך", "תיאור", "שעות", "תעריף שעתי (₪)", 'סה"כ (₪)']
    for cell, text in zip(header_cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    for entry in invoice.time_entries:
        line_total = entry.hours * entry.hourly_rate
        row_cells = table.add_row().cells
        row_cells[0].text = entry.entry_date.strftime("%d/%m/%Y")
        row_cells[1].text = entry.description
        row_cells[2].text = f"{entry.hours:g}"
        row_cells[3].text = f"{entry.hourly_rate:,.2f}"
        row_cells[4].text = f"{line_total:,.2f}"

    if not invoice.time_entries:
        empty_note = docx.add_paragraph()
        empty_note.add_run("לחשבונית זו אין רישומי שעות משויכים.").italic = True

    total_para = docx.add_paragraph()
    total_run = total_para.add_run(f'סה"כ לתשלום: {invoice.total_amount:,.2f} ₪')
    total_run.bold = True
    total_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if invoice.notes:
        note = docx.add_paragraph()
        note.add_run(f"הערות: {invoice.notes}").italic = True

    buffer = io.BytesIO()
    docx.save(buffer)
    return buffer.getvalue()

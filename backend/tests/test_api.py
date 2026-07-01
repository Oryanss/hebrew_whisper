def test_health(client):
    resp = client.get("/api/health")
    assert resp.status_code == 200
    assert resp.json() == {"status": "ok"}


def test_register_and_login(client):
    resp = client.post(
        "/api/auth/register",
        json={"full_name": "עו\"ד ראשון", "email": "lawyer1@example.com", "password": "Str0ngPass!23"},
    )
    assert resp.status_code == 201

    resp = client.post(
        "/api/auth/login", data={"username": "lawyer1@example.com", "password": "Str0ngPass!23"}
    )
    assert resp.status_code == 200
    assert "access_token" in resp.json()

    bad = client.post(
        "/api/auth/login", data={"username": "lawyer1@example.com", "password": "wrong"}
    )
    assert bad.status_code == 401


def test_endpoints_require_auth(client):
    resp = client.get("/api/cases")
    assert resp.status_code == 401


def test_client_and_case_lifecycle(client, auth_headers):
    resp = client.post(
        "/api/clients", json={"full_name": "ישראל ישראלי"}, headers=auth_headers
    )
    assert resp.status_code == 201
    client_id = resp.json()["id"]

    resp = client.post(
        "/api/cases",
        json={
            "case_number": "T-0001",
            "title": "תיק בדיקה",
            "practice_area": "חוזים",
            "client_id": client_id,
        },
        headers=auth_headers,
    )
    assert resp.status_code == 201
    case_id = resp.json()["id"]

    # duplicate case number should be rejected
    dup = client.post(
        "/api/cases",
        json={"case_number": "T-0001", "title": "כפילות", "client_id": client_id},
        headers=auth_headers,
    )
    assert dup.status_code == 400

    resp = client.get(f"/api/cases/{case_id}", headers=auth_headers)
    assert resp.status_code == 200
    assert resp.json()["title"] == "תיק בדיקה"


def test_draft_without_api_key_returns_503(client, auth_headers):
    client_resp = client.post(
        "/api/clients", json={"full_name": "לקוח לטיוטה"}, headers=auth_headers
    )
    client_id = client_resp.json()["id"]
    case_resp = client.post(
        "/api/cases",
        json={"case_number": "T-DRAFT-1", "title": "תיק לטיוטה", "client_id": client_id},
        headers=auth_headers,
    )
    case_id = case_resp.json()["id"]

    resp = client.post(
        f"/api/cases/{case_id}/draft",
        json={"doc_type": "מכתב התראה", "title": "מכתב", "instructions": "כתוב מכתב קצר"},
        headers=auth_headers,
    )
    assert resp.status_code == 503
    assert "ANTHROPIC_API_KEY" in resp.json()["detail"]


def test_citation_audit_missing_document_404s(client, auth_headers):
    resp = client.post("/api/documents/999999/audit-citations", headers=auth_headers)
    assert resp.status_code == 404


def test_citation_audit_flags_unverified_and_matches_verified(client, auth_headers):
    from app.database import SessionLocal
    from app import models

    client_resp = client.post(
        "/api/clients", json={"full_name": "לקוח לבדיקת אסמכתאות"}, headers=auth_headers
    )
    client_id = client_resp.json()["id"]
    case_resp = client.post(
        "/api/cases",
        json={"case_number": "T-AUDIT-1", "title": "תיק אסמכתאות", "client_id": client_id},
        headers=auth_headers,
    )
    case_id = case_resp.json()["id"]

    verified_citation = 'ע"א 1234/20'
    auth_resp = client.post(
        "/api/authorities",
        json={
            "citation_text": verified_citation,
            "source_type": "case_law",
            "source_database": "נבו",
            "case_id": case_id,
        },
        headers=auth_headers,
    )
    assert auth_resp.status_code == 201

    # Documents are normally only created via the /draft endpoint (which
    # requires a live LLM key); insert one directly here to exercise the
    # audit logic in isolation.
    db = SessionLocal()
    try:
        doc = models.Document(
            title="טיוטה לבדיקה",
            doc_type="מכתב",
            content=(
                'כפי שנקבע בע"א 1234/20 בפסק דין קודם, וכן ע"א 9999/99 שלא אומת מעולם.'
            ),
            case_id=case_id,
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)
        document_id = doc.id
    finally:
        db.close()

    resp = client.post(f"/api/documents/{document_id}/audit-citations", headers=auth_headers)
    assert resp.status_code == 200
    body = resp.json()
    citations = {f["citation_text"]: f["verified"] for f in body["findings"]}
    assert citations['ע"א 1234/20'] is True
    assert citations['ע"א 9999/99'] is False
    assert body["unverified_count"] == 1


def test_deadline_lifecycle_and_upcoming(client, auth_headers):
    from datetime import datetime, timedelta

    client_resp = client.post(
        "/api/clients", json={"full_name": "לקוח למועדים"}, headers=auth_headers
    )
    client_id = client_resp.json()["id"]
    case_resp = client.post(
        "/api/cases",
        json={"case_number": "T-DEADLINE-1", "title": "תיק מועדים", "client_id": client_id},
        headers=auth_headers,
    )
    case_id = case_resp.json()["id"]

    near_due = (datetime.utcnow() + timedelta(days=3)).isoformat()
    far_due = (datetime.utcnow() + timedelta(days=90)).isoformat()

    near = client.post(
        f"/api/cases/{case_id}/deadlines",
        json={"title": "הגשת כתב הגנה", "due_date": near_due},
        headers=auth_headers,
    )
    assert near.status_code == 201
    far = client.post(
        f"/api/cases/{case_id}/deadlines",
        json={"title": "דיון מקדמי", "due_date": far_due},
        headers=auth_headers,
    )
    assert far.status_code == 201

    listed = client.get(f"/api/cases/{case_id}/deadlines", headers=auth_headers)
    assert listed.status_code == 200
    assert len(listed.json()) == 2

    upcoming = client.get("/api/deadlines/upcoming?days=14", headers=auth_headers)
    assert upcoming.status_code == 200
    titles = [d["title"] for d in upcoming.json()]
    assert "הגשת כתב הגנה" in titles
    assert "דיון מקדמי" not in titles

    deadline_id = near.json()["id"]
    updated = client.patch(
        f"/api/deadlines/{deadline_id}", json={"status": "completed"}, headers=auth_headers
    )
    assert updated.status_code == 200
    assert updated.json()["status"] == "completed"

    upcoming_after = client.get("/api/deadlines/upcoming?days=14", headers=auth_headers)
    assert "הגשת כתב הגנה" not in [d["title"] for d in upcoming_after.json()]


def test_time_entries_and_billing_summary(client, auth_headers):
    client_resp = client.post(
        "/api/clients", json={"full_name": "לקוח לחיוב שעות"}, headers=auth_headers
    )
    client_id = client_resp.json()["id"]
    case_resp = client.post(
        "/api/cases",
        json={"case_number": "T-BILLING-1", "title": "תיק חיוב שעות", "client_id": client_id},
        headers=auth_headers,
    )
    case_id = case_resp.json()["id"]

    billable_with_rate = client.post(
        f"/api/cases/{case_id}/time-entries",
        json={
            "description": "ניסוח כתב תביעה",
            "entry_date": "2026-06-01T09:00:00",
            "hours": 3,
            "hourly_rate": 500,
            "billable": True,
        },
        headers=auth_headers,
    )
    assert billable_with_rate.status_code == 201

    billable_without_rate = client.post(
        f"/api/cases/{case_id}/time-entries",
        json={
            "description": "שיחת טלפון עם הלקוח",
            "entry_date": "2026-06-02T09:00:00",
            "hours": 1,
            "billable": True,
        },
        headers=auth_headers,
    )
    assert billable_without_rate.status_code == 201

    non_billable = client.post(
        f"/api/cases/{case_id}/time-entries",
        json={
            "description": "מחקר פנימי",
            "entry_date": "2026-06-03T09:00:00",
            "hours": 2,
            "hourly_rate": 500,
            "billable": False,
        },
        headers=auth_headers,
    )
    assert non_billable.status_code == 201

    listed = client.get(f"/api/cases/{case_id}/time-entries", headers=auth_headers)
    assert listed.status_code == 200
    assert len(listed.json()) == 3

    summary = client.get(f"/api/cases/{case_id}/billing-summary", headers=auth_headers)
    assert summary.status_code == 200
    body = summary.json()
    assert body["total_hours"] == 6
    assert body["billable_hours"] == 4
    assert body["total_billable_amount"] == 1500  # only the 3h @ 500 entry has a rate
    assert body["entries_missing_rate"] == 1  # the billable entry without a rate

    entry_id = billable_without_rate.json()["id"]
    patched = client.patch(
        f"/api/time-entries/{entry_id}", json={"hourly_rate": 400}, headers=auth_headers
    )
    assert patched.status_code == 200

    summary_after = client.get(f"/api/cases/{case_id}/billing-summary", headers=auth_headers)
    body_after = summary_after.json()
    assert body_after["total_billable_amount"] == 1900  # 1500 + 1h @ 400
    assert body_after["entries_missing_rate"] == 0


def test_case_notes_lifecycle(client, auth_headers):
    client_resp = client.post(
        "/api/clients", json={"full_name": "לקוח לרשומות"}, headers=auth_headers
    )
    client_id = client_resp.json()["id"]
    case_resp = client.post(
        "/api/cases",
        json={"case_number": "T-NOTES-1", "title": "תיק רשומות", "client_id": client_id},
        headers=auth_headers,
    )
    case_id = case_resp.json()["id"]

    empty = client.get(f"/api/cases/{case_id}/notes", headers=auth_headers)
    assert empty.status_code == 200
    assert empty.json() == []

    created = client.post(
        f"/api/cases/{case_id}/notes",
        json={"content": "שיחה עם הלקוח - מסכים להצעת הפשרה"},
        headers=auth_headers,
    )
    assert created.status_code == 201
    assert created.json()["created_by"]  # populated from the authenticated user
    assert created.json()["content"] == "שיחה עם הלקוח - מסכים להצעת הפשרה"

    listed = client.get(f"/api/cases/{case_id}/notes", headers=auth_headers)
    assert listed.status_code == 200
    assert len(listed.json()) == 1

    note_id = created.json()["id"]
    deleted = client.delete(f"/api/notes/{note_id}", headers=auth_headers)
    assert deleted.status_code == 204

    listed_after = client.get(f"/api/cases/{case_id}/notes", headers=auth_headers)
    assert listed_after.json() == []


def test_case_notes_missing_case_404s(client, auth_headers):
    resp = client.get("/api/cases/999999/notes", headers=auth_headers)
    assert resp.status_code == 404
    resp = client.post(
        "/api/cases/999999/notes", json={"content": "x"}, headers=auth_headers
    )
    assert resp.status_code == 404


def _make_docx_bytes(paragraphs):
    import io

    from docx import Document as DocxDocument

    docx = DocxDocument()
    for p in paragraphs:
        docx.add_paragraph(p)
    buffer = io.BytesIO()
    docx.save(buffer)
    return buffer.getvalue()


def test_upload_and_export_docx(client, auth_headers):
    client_resp = client.post(
        "/api/clients", json={"full_name": "לקוח למסמכי וורד"}, headers=auth_headers
    )
    client_id = client_resp.json()["id"]
    case_resp = client.post(
        "/api/cases",
        json={"case_number": "T-DOCX-1", "title": "תיק מסמכי וורד", "client_id": client_id},
        headers=auth_headers,
    )
    case_id = case_resp.json()["id"]

    docx_bytes = _make_docx_bytes(["שורה ראשונה", "שורה שנייה"])
    uploaded = client.post(
        f"/api/cases/{case_id}/documents/upload",
        data={"title": "כתב תביעה שהתקבל", "doc_type": "כתב תביעה"},
        files={
            "file": (
                "claim.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        headers=auth_headers,
    )
    assert uploaded.status_code == 201
    body = uploaded.json()
    assert "שורה ראשונה" in body["content"]
    assert "שורה שנייה" in body["content"]
    assert body["status"] == "in_review"

    # wrong extension is rejected
    rejected = client.post(
        f"/api/cases/{case_id}/documents/upload",
        data={"title": "קובץ שגוי", "doc_type": "אחר"},
        files={"file": ("claim.txt", b"not a docx", "text/plain")},
        headers=auth_headers,
    )
    assert rejected.status_code == 400

    document_id = body["id"]
    exported = client.get(f"/api/documents/{document_id}/export.docx", headers=auth_headers)
    assert exported.status_code == 200
    assert (
        exported.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert len(exported.content) > 0


def test_knowledge_library_upload_search_and_delete(client, auth_headers):
    txt_upload = client.post(
        "/api/knowledge/upload",
        data={"title": "מאמר על התיישנות", "category": "article"},
        files={"file": ("article.txt", "דיון מקיף בסוגיית ההתיישנות בדין הישראלי".encode("utf-8"), "text/plain")},
        headers=auth_headers,
    )
    assert txt_upload.status_code == 201
    doc_id = txt_upload.json()["id"]
    assert "content" not in txt_upload.json()  # list/create responses omit full content

    docx_bytes = _make_docx_bytes(["פסק דין לדוגמה", "נקבע כי יש לדחות את הערעור"])
    docx_upload = client.post(
        "/api/knowledge/upload",
        data={"title": "פסק דין לדוגמה", "category": "case_law"},
        files={
            "file": (
                "example.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        headers=auth_headers,
    )
    assert docx_upload.status_code == 201

    unsupported = client.post(
        "/api/knowledge/upload",
        data={"title": "קובץ לא נתמך", "category": "other"},
        files={"file": ("image.png", b"\x89PNG", "image/png")},
        headers=auth_headers,
    )
    assert unsupported.status_code == 400

    listed = client.get("/api/knowledge", headers=auth_headers)
    assert listed.status_code == 200
    assert len(listed.json()) >= 2

    filtered = client.get("/api/knowledge?category=case_law", headers=auth_headers)
    assert all(d["category"] == "case_law" for d in filtered.json())

    searched = client.get("/api/knowledge/search?q=התיישנות", headers=auth_headers)
    assert searched.status_code == 200
    assert any("התיישנות" in r["snippet"] for r in searched.json())

    detail = client.get(f"/api/knowledge/{doc_id}", headers=auth_headers)
    assert detail.status_code == 200
    assert "content" in detail.json()

    deleted = client.delete(f"/api/knowledge/{doc_id}", headers=auth_headers)
    assert deleted.status_code == 204
    missing = client.get(f"/api/knowledge/{doc_id}", headers=auth_headers)
    assert missing.status_code == 404


def test_legal_research_without_api_key_returns_503(client, auth_headers):
    resp = client.post(
        "/api/legal-research",
        json={"query": "מהי ההלכה לגבי התיישנות בתביעת חוב?"},
        headers=auth_headers,
    )
    assert resp.status_code == 503
    assert "ANTHROPIC_API_KEY" in resp.json()["detail"]


def test_risk_assessment_scoring_and_lifecycle(client, auth_headers):
    client_resp = client.post(
        "/api/clients", json={"full_name": "לקוח להערכת סיכון"}, headers=auth_headers
    )
    client_id = client_resp.json()["id"]
    case_resp = client.post(
        "/api/cases",
        json={"case_number": "T-RISK-1", "title": "תיק הערכת סיכון", "client_id": client_id},
        headers=auth_headers,
    )
    case_id = case_resp.json()["id"]

    low = client.post(
        f"/api/cases/{case_id}/risk-assessments",
        json={
            "category": "contract",
            "description": "סטייה קלה מנוסח סטנדרטי בחוזה ספק",
            "severity": 1,
            "likelihood": 2,
        },
        headers=auth_headers,
    )
    assert low.status_code == 201
    low_body = low.json()
    assert low_body["risk_score"] == 2
    assert low_body["risk_level"] == "green"
    assert low_body["assessed_by"]

    critical = client.post(
        f"/api/cases/{case_id}/risk-assessments",
        json={
            "category": "litigation",
            "description": "תביעה פעילה עם חשיפה כספית משמעותית",
            "severity": 5,
            "likelihood": 4,
        },
        headers=auth_headers,
    )
    assert critical.status_code == 201
    critical_body = critical.json()
    assert critical_body["risk_score"] == 20
    assert critical_body["risk_level"] == "red"
    assert "הסלמה מיידית" in critical_body["recommended_action"]

    listed = client.get(f"/api/cases/{case_id}/risk-assessments", headers=auth_headers)
    assert listed.status_code == 200
    assert len(listed.json()) == 2
    # newest first
    assert listed.json()[0]["id"] == critical_body["id"]

    invalid = client.post(
        f"/api/cases/{case_id}/risk-assessments",
        json={"description": "ציון לא חוקי", "severity": 6, "likelihood": 3},
        headers=auth_headers,
    )
    assert invalid.status_code == 422

    deleted = client.delete(
        f"/api/risk-assessments/{low_body['id']}", headers=auth_headers
    )
    assert deleted.status_code == 204
    listed_after = client.get(f"/api/cases/{case_id}/risk-assessments", headers=auth_headers)
    assert len(listed_after.json()) == 1


def test_invoice_docx_generation(client, auth_headers):
    import io

    from docx import Document as DocxDocument

    client_resp = client.post(
        "/api/clients", json={"full_name": "לקוח לחשבונית"}, headers=auth_headers
    )
    client_id = client_resp.json()["id"]
    # Hebrew case title is exactly the scenario that broke the original
    # document-export endpoint before the RFC 5987 filename fix - keep it.
    case_resp = client.post(
        "/api/cases",
        json={
            "case_number": "T-INVOICE-1",
            "title": "תיק חשבונית לבדיקה - חברת בע\"מ",
            "client_id": client_id,
        },
        headers=auth_headers,
    )
    case_id = case_resp.json()["id"]

    billable_with_rate = client.post(
        f"/api/cases/{case_id}/time-entries",
        json={
            "description": "ניסוח כתב הגנה",
            "entry_date": "2026-06-01T09:00:00",
            "hours": 4,
            "hourly_rate": 600,
            "billable": True,
        },
        headers=auth_headers,
    )
    assert billable_with_rate.status_code == 201

    billable_without_rate = client.post(
        f"/api/cases/{case_id}/time-entries",
        json={
            "description": "ישיבת ייעוץ עם הלקוח",
            "entry_date": "2026-06-02T09:00:00",
            "hours": 1.5,
            "billable": True,
        },
        headers=auth_headers,
    )
    assert billable_without_rate.status_code == 201

    non_billable = client.post(
        f"/api/cases/{case_id}/time-entries",
        json={
            "description": "מחקר פנימי שאינו לחיוב",
            "entry_date": "2026-06-03T09:00:00",
            "hours": 2,
            "hourly_rate": 600,
            "billable": False,
        },
        headers=auth_headers,
    )
    assert non_billable.status_code == 201

    invoice = client.get(f"/api/cases/{case_id}/invoice.docx", headers=auth_headers)
    assert invoice.status_code == 200
    assert (
        invoice.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert len(invoice.content) > 0

    # This is exactly the class of bug that previously crashed with
    # UnicodeEncodeError: Content-Disposition must be RFC 5987 encoded when
    # the filename contains Hebrew, plus an ASCII-only fallback.
    content_disposition = invoice.headers["content-disposition"]
    assert "filename*=UTF-8''" in content_disposition
    assert 'filename="' in content_disposition
    fallback_match = content_disposition.split('filename="', 1)[1].split('"', 1)[0]
    assert fallback_match.isascii()

    # Verify the itemized content of the generated docx: only the
    # billable + priced entry should appear, the total should be correct,
    # and the excluded-for-missing-rate count should be noted.
    docx = DocxDocument(io.BytesIO(invoice.content))
    full_text = "\n".join(p.text for p in docx.paragraphs)
    assert "לקוח לחשבונית" in full_text
    assert "תיק חשבונית לבדיקה" in full_text
    assert "1" in full_text  # entries_missing_rate note mentions the excluded entry

    table_text = []
    for table in docx.tables:
        for row in table.rows:
            table_text.append("\t".join(cell.text for cell in row.cells))
    joined_table = "\n".join(table_text)
    assert "ניסוח כתב הגנה" in joined_table
    assert "ישיבת ייעוץ עם הלקוח" not in joined_table  # missing rate, excluded
    assert "מחקר פנימי שאינו לחיוב" not in joined_table  # not billable, excluded
    assert "2,400.00" in joined_table  # 4h * 600 line total
    assert "2,400.00" in full_text  # total line


def test_invoice_docx_with_no_time_entries(client, auth_headers):
    import io

    from docx import Document as DocxDocument

    client_resp = client.post(
        "/api/clients", json={"full_name": "לקוח ללא רישומי שעות"}, headers=auth_headers
    )
    client_id = client_resp.json()["id"]
    case_resp = client.post(
        "/api/cases",
        json={"case_number": "T-INVOICE-EMPTY-1", "title": "תיק ריק", "client_id": client_id},
        headers=auth_headers,
    )
    case_id = case_resp.json()["id"]

    invoice = client.get(f"/api/cases/{case_id}/invoice.docx", headers=auth_headers)
    assert invoice.status_code == 200
    assert (
        invoice.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert len(invoice.content) > 0

    docx = DocxDocument(io.BytesIO(invoice.content))
    full_text = "\n".join(p.text for p in docx.paragraphs)
    assert "אין רישומי שעות לחיוב" in full_text


def test_invoice_docx_missing_case_returns_404(client, auth_headers):
    resp = client.get("/api/cases/999999/invoice.docx", headers=auth_headers)
    assert resp.status_code == 404
